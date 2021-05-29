# 해야하는 것
# 1. word파일 생성 -- O
# 2. txt파일 읽고 그 내용 word파일에 입력
# 3. 이미지를 word파일에 삽입 -- 0
# 4. word파일의 폰트, 기울기, size , bold 등 적용 가능 -- 0
# 5. word파일을 원하는 경로에 저장 -- 0 

import datetime
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH # 가운데 정렬 모듈
from docx.oxml.ns import qn # 한글 출력 모듈
from docx.shared import Pt, RGBColor, Inches# 색 변경 모듈
from docx.oxml.ns import qn # 한글 폰트

# ------------ 시간 출력
now = datetime.datetime.now()
nowDate = now.strftime('%Y-%m-%d')
nowTime = now.strftime('%H:%M:%S')
# -------------------- 

# word 파일 생성
doc = docx.Document()

# title 생성
doc.add_heading(' Event Shop\'s Report ', level=0)

# 문단 추가
pa0 = doc.add_paragraph("Date :  " + nowDate + " Time : " + nowTime)
pa1 = doc.add_paragraph("\n\n\n 1. paragraph : EC2 : 스타일: Intense Quote", style= 'Intense Quote')
pa2 = doc.add_paragraph("\n\n")
run2 = pa2.add_run('2. paragraph : RDS.')
run2.font.size = docx.shared.Pt(20) # 글씨 크기 20으로

pa3 = doc.add_paragraph("\n\n")
run3 = pa3.add_run(' 3. paragraph : S3.')
run3.font.size = docx.shared.Pt(20) # 글씨 크기 20으로

pa4 = doc.add_paragraph("\n\n")
run4 = pa4.add_run(' 4. paragraph : Lambda.')
run4.font.size = docx.shared.Pt(20) # 글씨 크기 20으로

# pont 적용
pa0.alignment = WD_ALIGN_PARAGRAPH.RIGHT # 오른쪽 정렬
# pa0.alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬
# pa0.alignment = WD_ALIGN_PARAGRAPH.LEFT # 왼쪽 정렬
# pa1.font.name = 'Arial'
pa1.bold = True
# pa1.font.szie = docx.shared.Pt(20)

pa2.italic = True # 기울기


# 4. 문단에 글 추가
run4_1 = pa4.add_run('\n 내용 추가 : 글씨 크기 20 ')
run4_2 = pa4.add_run('\n' + nowDate)
run4_3 = pa4.add_run('\n 글씨 추가 : 글씨 크기 15 ')
run4_4 = pa4.add_run('\n 파란색 글씨')

# run 스타일 적용
run4.font.name = 'Arial' # 글씨체 'Arial'로 수정
run4.font.size = docx.shared.Pt(20) # 글씨 크기 20으로
run4_3.font.size = docx.shared.Pt(15) # 글씨 크기 15으로
run4_2.italic = True
run4_3.bold = True
run4_3.font.color.rgb = RGBColor(255,0,0) # 빨간색
run4_4.font.color.rgb = RGBColor(0,0,255) # 파란색

# run4_4._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조') # 한글 폰트 적용 안됨;;

# 페이지 넘어가기
doc.add_page_break()
doc.add_heading(' 2. Page ', level=0)
pa6 = doc.add_paragraph("6. cat ")

# 5. 사진 추가
doc.add_picture('cat.png', width=docx.shared.Cm(5), height = docx.shared.Cm(5))
pic2 = doc.add_picture('cat2.png', width = docx.shared.Cm(7), height = docx.shared.Cm(7))
pic2.alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬

# records라는 테이블을 추가
pa5 = doc.add_paragraph("\n 5. paragraph : graph. ")

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = doc.add_table(rows=1, cols=3)
table.style='Colorful Shading Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ec2'
hdr_cells[1].text = 'cpu'
hdr_cells[2].text = 'Number'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
	
hdr_cells[0].width = Inches(2.9)
hdr_cells[0].height = Inches(2.9)

# 3 페이지 넘어가기
doc.add_page_break()
doc.add_heading(' 3. Page ', level=0)

# 4 페이지 넘어가기
doc.add_page_break()
doc.add_heading(' 4. Page ', level=0)

# beom이라는 이름으로 docx 저장
doc.save('C:/Users/dolif/Desktop/프로젝트/word/beom.docx')

print("success")
